# ******************************************************************************
#
# vocutil, educational vocabulary utilities.
#
# Copyright 2022-2024 Jeremy A Gray <gray@flyquackswim.com>.
#
# All rights reserved.
#
# SPDX-License-Identifier: GPL-3.0-or-later
#
# ******************************************************************************

"""vocutil testgen tests."""

import pytest
from docx import Document
from docx.api import _default_docx_path

import vocutil


def test__parse_testgen_mc_parseable(question_data):
    """Should parse TestGen multiple choice question banks."""
    q = vocutil._parse_testgen_mc(question_data["text"])
    print(q)

    assert question_data["question"] == q[0]["question"]
    assert question_data["chapter"] == q[0]["chapter"]
    assert question_data["section"] == q[0]["section"]
    assert q[0]["answers"][question_data["answer"]]["correct"]

    # Wrong answers.
    for i, answer in enumerate(q[0]["answers"]):
        if i != question_data["answer"]:
            assert not q[0]["answers"][i]["correct"]


def test__parse_testgen_mc_unparseable(unparseable):
    """Should parse TestGen multiple choice question banks."""
    assert [] == vocutil._parse_testgen_mc(unparseable)


def test__load_testgen_file_valid_docx(fs, question_data):
    """Should load a MS DOCX file."""
    # python-docx loads a default template from the real file system
    # that is not available to the fake file system unless it is
    # added.
    fs.add_real_file(_default_docx_path())

    fn = "./testbank.docx"
    fs.create_file(fn)

    doc = Document()
    doc.add_paragraph(question_data["text"])
    doc.save(fn)

    assert question_data["text"] == vocutil._load_testgen_file(fn)


def test__load_testgen_file_valid_text(fs, question_data):
    """Should load a text file."""
    fn = "./testbank.txt"
    fs.create_file(fn)
    with open(fn, "w") as f:
        f.write(question_data["text"])

    assert question_data["text"] == vocutil._load_testgen_file(fn)


def test__load_testgen_file_bad_filename(fs, badfn):
    """Should raise ``FileNotFoundError`` on bad file names."""
    with pytest.raises(FileNotFoundError):
        vocutil._load_testgen_file(badfn)


def test_load_testgen_mc(fs, question_data):
    """Should load and parse a valid file."""
    # python-docx loads a default template from the real file system
    # that is not available to the fake file system unless it is
    # added.
    fs.add_real_file(_default_docx_path())

    fn = "./testbank.docx"
    fs.create_file(fn)

    doc = Document()
    doc.add_paragraph(question_data["text"])
    doc.save(fn)

    q = vocutil.load_testgen_mc(fn)

    assert question_data["question"] == q[0]["question"]
    assert question_data["chapter"] == q[0]["chapter"]
    assert question_data["section"] == q[0]["section"]
    assert q[0]["answers"][question_data["answer"]]["correct"]

    # Wrong answers.
    for i, answer in enumerate(q[0]["answers"]):
        if i != question_data["answer"]:
            assert not q[0]["answers"][i]["correct"]
